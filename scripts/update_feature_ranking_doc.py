from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Classement_actualise_fonctionnalites_BotRoad.docx"


ACCENT = RGBColor(13, 93, 122)
GREEN = RGBColor(46, 125, 50)
ORANGE = RGBColor(196, 117, 0)
RED = RGBColor(198, 40, 40)
GRAY = RGBColor(80, 80, 80)


def set_run(run, size=10, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Classement actualise des fonctionnalites - BotRoad")
    set_run(run, size=20, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Etat du projet apres integration des alertes, statistiques et mode itineraire sur")
    set_run(run, size=10, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mise a jour : 02/06/2026")
    set_run(run, size=9, color=GRAY)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        run.text = text
    else:
        p.runs[0].text = text
    set_run(p.runs[0], size=14 if level == 1 else 12, bold=True, color=ACCENT)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16.0)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shading = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "E7F6FB")
    shading.append(shd)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run(run, size=10, bold=True, color=ACCENT)
    p.add_run("\n")
    run = p.add_run(body)
    set_run(run, size=9, color=GRAY)
    doc.add_paragraph()


def status_color(status):
    if status == "FAIT":
        return GREEN
    if status == "PRIORITAIRE":
        return RED
    if status == "EN COURS":
        return ORANGE
    return GRAY


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [Cm(5.2), Cm(2.7), Cm(2.9), Cm(6.0)]
    headers = ["Fonctionnalite", "Statut", "Priorite", "Commentaire"]

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run(run, size=9, bold=True, color=RGBColor(255, 255, 255))
        shade_cell(cell, "0D5D7A")

    for item in rows:
        cells = table.add_row().cells
        values = [item["name"], item["status"], item["priority"], item["comment"]]
        for i, value in enumerate(values):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [1, 2] else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            color = status_color(item["status"]) if i == 1 else RGBColor(30, 30, 30)
            set_run(run, size=8.5, bold=(i == 1), color=color)
    doc.add_paragraph()


def shade_cell(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def p(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    add_title(doc)
    add_note(
        doc,
        "Resume rapide",
        "Le coeur du projet est maintenant bien avance : chatbot, calcul d'itineraire, carte, signalements, alertes, statistiques, historique par trajet et mode itineraire sur sont fonctionnels. La suite logique concerne surtout les tests reels, la partie IoT et le futur modele de classification.",
    )

    add_heading(doc, "1. Fonctionnalites deja realisees")
    done_rows = [
        {"name": "Chatbot texte", "status": "FAIT", "priority": "Coeur MVP", "comment": "Dialogue utilisateur et traitement des demandes simples."},
        {"name": "Calcul d'itineraire", "status": "FAIT", "priority": "Coeur MVP", "comment": "Calcul via Google Directions avec distance, duree et geometrie."},
        {"name": "Carte interactive", "status": "FAIT", "priority": "Coeur MVP", "comment": "Affichage carte, zoom, navigation et reouverture apres fermeture."},
        {"name": "Instructions par segments", "status": "FAIT", "priority": "Important", "comment": "Affichage des etapes de navigation et segments dangereux."},
        {"name": "Alertes sur la carte", "status": "FAIT", "priority": "Coeur MVP", "comment": "Marqueurs et zones a risque rattaches aux signalements."},
        {"name": "Signalement manuel", "status": "FAIT", "priority": "Coeur MVP", "comment": "Bouton Signaler, position GPS, type, gravite et commentaire."},
        {"name": "Signalement par chatbot", "status": "FAIT", "priority": "Important", "comment": "L'utilisateur peut declarer une mauvaise route par message."},
        {"name": "Validite 48h", "status": "FAIT", "priority": "Important", "comment": "Une alerte expire automatiquement apres 48h sauf suppression avant."},
        {"name": "Confirmation/refutation", "status": "FAIT", "priority": "Important", "comment": "Systeme type like/dislike pour renforcer ou contester une alerte."},
        {"name": "Page Alertes", "status": "FAIT", "priority": "Important", "comment": "Liste, details, filtres, statut, prise en charge et suppression admin."},
        {"name": "Statistiques simples", "status": "FAIT", "priority": "Bonus utile", "comment": "Total, actives, expirees, gravite dominante, confirmations/refutations."},
        {"name": "Historique alertes par trajet", "status": "FAIT", "priority": "Important", "comment": "Chaque trajet affiche les alertes detectees au moment du calcul."},
        {"name": "Itineraire alternatif", "status": "FAIT", "priority": "Coeur securite", "comment": "Recherche d'une alternative lorsque la route contient des alertes."},
        {"name": "Mode itineraire sur", "status": "FAIT", "priority": "Coeur securite", "comment": "Choix rapide/sur/equilibre selon duree, distance et score de risque."},
    ]
    add_table(doc, done_rows)

    add_heading(doc, "2. Fonctionnalites primordiales restantes")
    priority_rows = [
        {"name": "Tests reels multi-appareils", "status": "PRIORITAIRE", "priority": "Tres haute", "comment": "Verifier signalements, votes, alertes et alternatives entre deux utilisateurs."},
        {"name": "Regles Firestore et index", "status": "PRIORITAIRE", "priority": "Tres haute", "comment": "Securiser l'acces aux donnees et eviter les erreurs de requetes en production."},
        {"name": "Expiration cote serveur", "status": "PRIORITAIRE", "priority": "Haute", "comment": "Ajouter une tache ou logique serveur pour nettoyer/archiver les alertes expirees."},
        {"name": "Workflow administrateur", "status": "EN COURS", "priority": "Haute", "comment": "La prise en charge existe ; il faut tester et stabiliser le role admin."},
        {"name": "Alertes automatiques IoT", "status": "PRIORITAIRE", "priority": "Memoire", "comment": "Integrer les donnees capteurs pour detecter trous, chocs ou mauvaises routes."},
        {"name": "Modele DL de classification", "status": "PRIORITAIRE", "priority": "Memoire", "comment": "Classifier les types de mauvaises routes ; le score actuel est pret a l'accueillir."},
        {"name": "Evaluation du score de risque", "status": "PRIORITAIRE", "priority": "Memoire", "comment": "Comparer score regles + futures predictions DL sur plusieurs trajets."},
    ]
    add_table(doc, priority_rows)

    add_heading(doc, "3. Fonctionnalites moins urgentes ou bonus")
    bonus_rows = [
        {"name": "Entree vocale", "status": "A VENIR", "priority": "Bonus", "comment": "Conversion voix vers texte pour commander le chatbot."},
        {"name": "Synthese vocale", "status": "A VENIR", "priority": "Bonus", "comment": "Lire les reponses du chatbot et les alertes de navigation."},
        {"name": "Guidage vocal", "status": "A VENIR", "priority": "Bonus", "comment": "Instructions vocales type tourner a gauche, continuer tout droit."},
        {"name": "Lieux proches", "status": "PARTIEL", "priority": "Secondaire", "comment": "Recherche de lieux proche deja presente cote chatbot, a renforcer cote carte."},
        {"name": "Filtres lieux par categorie", "status": "A VENIR", "priority": "Secondaire", "comment": "Utile mais moins lie au coeur des routes dangereuses."},
        {"name": "Statistiques avancees", "status": "A VENIR", "priority": "Bonus", "comment": "Graphiques, tendances, zones les plus signalees, evolution temporelle."},
        {"name": "Amelioration UI carte", "status": "A VENIR", "priority": "Confort", "comment": "Rendre les segments et alertes encore plus lisibles visuellement."},
    ]
    add_table(doc, bonus_rows)

    add_heading(doc, "4. Position actuelle du projet")
    p(doc, "Le MVP applicatif est suffisamment avance pour commencer les tests reels. La partie application mobile couvre deja l'interaction utilisateur, la carte, le signalement, la consultation des alertes, l'historique et la proposition d'itineraire plus sur.")
    p(doc, "La prochaine grande etape technique est de connecter la partie IoT et de preparer le modele de deep learning qui classifiera les types de mauvaises routes. Le score de risque actuel doit etre considere comme une base provisoire : il exploite les alertes utilisateur, la gravite et la validation collaborative, mais il pourra etre enrichi par les predictions du modele.")

    add_heading(doc, "5. Prochaine etape conseillee")
    next_rows = [
        {"name": "1. Tester le parcours complet", "status": "PRIORITAIRE", "priority": "Maintenant", "comment": "Creer trajet, signaler, confirmer/refuter, verifier alerte sur autre appareil."},
        {"name": "2. Stabiliser backend", "status": "PRIORITAIRE", "priority": "Maintenant", "comment": "Regles Firestore, index, expiration, roles utilisateur/admin."},
        {"name": "3. Preparer IoT", "status": "PRIORITAIRE", "priority": "Suite memoire", "comment": "Definir donnees capteur, format d'envoi et stockage."},
        {"name": "4. Integrer classification DL", "status": "A VENIR", "priority": "Suite memoire", "comment": "Classifier les anomalies puis alimenter le score de risque."},
    ]
    add_table(doc, next_rows)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
